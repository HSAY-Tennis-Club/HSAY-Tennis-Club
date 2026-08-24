from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_hsay_ui_doc import (
    CORAL,
    GREEN,
    INK,
    LIME,
    MUTED,
    PAPER,
    RED,
    ROOT,
    VIOLET,
    WHITE,
    add_callout,
    add_h1,
    add_h2,
    add_kicker,
    add_label_detail,
    add_matrix,
    add_page_break,
    add_para,
    set_para,
    set_run,
    set_table_geometry,
    setup_document,
    shade,
)


OUTPUT = ROOT / "docs" / "HSAY_VI品牌视觉识别规范_v1.0.docx"


def build():
    doc = setup_document()
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    set_para(header, 0, 0, 1.0)
    set_run(header.add_run("HSAY  /  VISUAL IDENTITY GUIDELINES"), 8, MUTED, True)

    # 01 — Cover
    add_para(doc, "HSAY · HIT / SPIN / ACE / YOU", 9, CORAL, True, after=70, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "HSAY", 36, INK, True, after=3, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    add_para(doc, "品牌视觉识别系统 VI 规范", 22, INK, True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.05)
    add_para(doc, "场下是宝贝，场上撕得飞。", 14, VIOLET, True, after=34, align=WD_ALIGN_PARAGRAPH.CENTER)
    brand = doc.add_table(rows=1, cols=6)
    for cell, color in zip(brand.rows[0].cells, [CORAL, "FF9C2A", LIME, "44DFAC", "4DA6FF", VIOLET]):
        shade(cell, color)
        cell.text = ""
    set_table_geometry(brand, [1560] * 6, indent=0)
    add_para(doc, "", after=28)
    add_callout(doc, "VERSION 1.0 · 2026.08.24", "硬核竞技 × 社群温度 × 无可救药的抓马", "本规范定义 HSAY 的品牌核心、Logo 方向、色彩、字体、图形语言、影像、数据视觉和落地应用。它不是界面说明书，而是所有视觉触点共同遵守的品牌识别母版。", PAPER, VIOLET, INK)
    add_para(doc, "适用：网站、小程序、赛事海报、社交媒体、赛果卡、球员视觉、服装与俱乐部周边。", 9.5, MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 02 — Brand core
    add_page_break(doc)
    add_kicker(doc, "01 · BRAND CORE")
    add_h1(doc, "品牌内核与性格坐标")
    add_para(doc, "HSAY 扎根上海，服务兼具竞技实力和情绪价值的网球社群。品牌必须同时让人感到专业、包容、自信与好玩：场下互相托住，场上认真分胜负。")
    add_callout(doc, "BRAND PROMISE", "场下是宝贝，场上撕得飞。", "这是对内的关系准则，也是对外最稳定的品牌记忆。任何表达都应保留竞技锋利感，同时避免攻击、排斥或刻板化。")
    add_h2(doc, "四个品牌支柱")
    add_matrix(doc, ["字母", "英文", "品牌含义", "视觉联想"], [
        ("H", "HIT", "爆抽、力量、上场不退让", "粗重切面、强对比、直线冲击"),
        ("S", "SPIN", "控制、旋转、体面地“纯属我演”", "弧线轨迹、错位、动态倾斜"),
        ("A", "ACE", "高光得分与自信闪耀", "荧光绿、星点、瞬时爆发"),
        ("Y", "YOU", "缺你不可的社群连接", "双向关系、并列名字、彩虹连接"),
    ], [900, 1200, 3900, 3360])
    add_h2(doc, "性格尺度")
    add_label_detail(doc, [
        ("专业 70 / 亲近 30", "比赛信息清晰可信；社群语言可以俏皮，但不能模糊事实。"),
        ("锋利 65 / 柔软 35", "标题和构图有冲击力；涉及成员时保持尊重与安全感。"),
        ("张扬 60 / 克制 40", "重点允许荧光和彩虹爆发；信息密集区必须留白。"),
        ("包容 100", "LGBTQ+ 友好通过语言、影像与真实关系自然表达，不做标签消费。"),
    ])

    # 03 — Name and mark
    add_page_break(doc)
    add_kicker(doc, "02 · NAME & MARK")
    add_h1(doc, "名称、字标与安全空间")
    add_h2(doc, "标准名称")
    add_matrix(doc, ["场景", "推荐写法", "说明"], [
        ("正式英文", "HSAY", "全大写；作为主字标和赛事署名"),
        ("英文释义", "Hit · Spin · Ace & You", "首次完整介绍时使用，不替代主字标"),
        ("中文昵称", "沪上阿姨", "用于社群语境；不与其他商业品牌的识别混用"),
        ("组合署名", "HSAY 沪上阿姨网球俱乐部", "适用于活动报名、规则与合作物料"),
    ], [1600, 3400, 4360])
    add_h2(doc, "HSAY 字标构造方向")
    add_callout(doc, "PRIMARY WORDMARK", "HSAY", "使用现代极粗无衬线骨架，字距略紧；允许在字标下方加入一条六色 Spin 轨迹。轨迹必须保持单向流动感，不能画成封闭彩虹圈或装饰泡泡。", INK, LIME, WHITE)
    add_h2(doc, "安全空间与最小尺寸")
    add_label_detail(doc, [
        ("安全空间", "以字标 H 的竖画宽度 X 为基准，四周至少保留 1X。"),
        ("数字最小宽度", "完整字标不小于 72px；仅字母标不小于 28px。"),
        ("印刷最小宽度", "完整字标不小于 24mm；刺绣应用不小于 32mm。"),
        ("背景", "优先 Ink、Paper 或纯色；复杂照片上必须增加纯色承托区。"),
        ("禁止", "不得拉伸、描边、渐变字面、任意旋转、改变字母次序或叠加球拍与爱心。"),
    ])

    # 04 — Color
    add_page_break(doc)
    add_kicker(doc, "03 · COLOR SYSTEM")
    add_h1(doc, "品牌色彩与配比")
    palette = doc.add_table(rows=1, cols=6)
    for cell, color in zip(palette.rows[0].cells, [INK, PAPER, LIME, VIOLET, CORAL, GREEN]):
        shade(cell, color)
        cell.text = ""
    set_table_geometry(palette, [1560] * 6, indent=0)
    add_para(doc, "", after=10)
    add_matrix(doc, ["角色", "HEX", "占比", "用途", "限制"], [
        ("Ink", "#141414", "35%", "字标、标题、边框、夜场物料", "避免整页纯黑压迫"),
        ("Paper", "#F4F1EA", "35%", "主背景、留白、纸张感", "避免脏灰和低对比"),
        ("Ace Lime", "#D7FF38", "12%", "高光、CTA、冠军与 Ace", "不可用于长段正文"),
        ("Spin Violet", "#8263FF", "8%", "旋转、阴影、会员画像", "白字组合需检查对比"),
        ("Drama Coral", "#FF654F", "6%", "抓马强调、眉题、活动热度", "不替代负方语义色"),
        ("Court Green", "#087958", "4%", "胜方、球场、稳定状态", "必须配合 WIN / W"),
    ], [1400, 1300, 900, 3260, 2500])
    add_h2(doc, "六色 Pride 轨迹")
    add_para(doc, "彩虹采用 Coral → Orange → Ace Lime → Mint → Blue → Violet 的固定顺序。优先作为 4–10px 的运动轨迹、底边或切面，而非大面积渐变背景。每个画面最多出现一条主轨迹。")
    add_h2(doc, "胜负语义色")
    add_label_detail(doc, [
        ("WIN / W", "Court Green #087958；同时显示文字或图标。"),
        ("LOST / L", "Brick Red #DC5D49；不与 Drama Coral 混用。"),
        ("中性", "Muted Gray #716E68；用于平局、未开始、无变化。"),
    ])

    # 05 — Type and graphics
    add_page_break(doc)
    add_kicker(doc, "04 · TYPE & GRAPHICS")
    add_h1(doc, "字体层级与动态图形")
    add_h2(doc, "字体系统")
    add_label_detail(doc, [
        ("中文展示", "PingFang SC Heavy / Microsoft YaHei Bold；超粗、紧字距、短句。"),
        ("英文展示", "Geist Black / Arial Black；用于 HSAY、比分与赛事标题。"),
        ("正文", "PingFang SC Regular / Microsoft YaHei；保持 1.5–1.8 倍行距。"),
        ("数据与标签", "Geist Mono / 等宽字体；用于日期、比分、排名、H/S/A/Y。"),
    ])
    add_h2(doc, "排版气质")
    add_matrix(doc, ["层级", "建议", "示例"], [
        ("主标题", "超粗、2–3 行以内、左对齐优先", "场下是宝贝，场上撕得飞。"),
        ("数据", "数字比单位至少大 1.5 倍", "6–3 / #01 / 2026.08.12"),
        ("眉题", "全大写、小字号、高字距", "TEAM FINAL / ACE MOMENT"),
        ("正文", "自然中文，不做过度英文拼贴", "查赛程、看赛果、追排名。"),
    ], [1400, 4000, 3960])
    add_h2(doc, "Spin 轨迹与切面")
    add_callout(doc, "SIGNATURE DEVICE", "一道轨迹，穿过每一次高光", "核心动态图形是具有方向性的六色弧线：从画面外进入，擦过字标或关键数字，再离开画面。搭配 1px 直线、硬阴影和轻微 1–3° 倾斜，形成网球速度感。除头像与网球本体外，避免大量圆形相互叠压。", VIOLET, LIME, WHITE)

    # 06 — Imagery
    add_page_break(doc)
    add_kicker(doc, "05 · IMAGE DIRECTION")
    add_h1(doc, "影像、人物与社群表达")
    add_h2(doc, "摄影关键词")
    add_matrix(doc, ["方向", "要", "不要"], [
        ("竞技", "击球瞬间、脚步、汗水、网线与球印", "只有摆拍握拍和空球场"),
        ("关系", "击掌、场边互动、双打默契、赛后拥抱", "把成员处理成孤立模特"),
        ("上海", "夜场灯光、城市肌理、真实俱乐部场地", "泛城市天际线素材拼贴"),
        ("Pride", "真实、多样、自然的成员呈现", "符号堆砌、彩虹滤镜覆盖人脸"),
        ("抓马", "胜负表情、庆祝、幽默瞬间", "恶意丑化、羞辱或过度猎奇"),
    ], [1300, 4040, 4020])
    add_h2(doc, "裁切与后期")
    add_label_detail(doc, [
        ("裁切", "人物可贴边、动作可越框；球拍与球的方向要保留运动空间。"),
        ("色调", "中性肤色 + 略高对比；夜场允许深黑和 Ace Lime 点亮。"),
        ("叠字", "优先放在低细节区域；照片上叠字必须使用 Paper 或 Ink 承托。"),
        ("头像", "球员无照片时使用单字母 / 单汉字色块头像；同一球员始终同色。"),
    ])
    add_h2(doc, "隐私与尊重")
    add_para(doc, "公开物料只使用获得授权的照片与公开赛果。密友备注、私密训练评价、联系方式和身份信息不得进入海报、社交媒体或公开球员卡。")

    # 07 — Score and data expression
    add_page_break(doc)
    add_kicker(doc, "06 · SCORE & DATA EXPRESSION")
    add_h1(doc, "比分、排名与赛果主视觉")
    add_h2(doc, "赛果主视觉骨架")
    add_label_detail(doc, [
        ("顶部", "赛事类型 + 日期；小字号等宽，建立可信上下文。"),
        ("中部", "双方名称、标识与总比分；比分是第一视觉焦点。"),
        ("底部", "冠军结论或一句品牌化标题；避免装饰抢过数据。"),
        ("背景", "Ink / Paper 单色底 + 球场分割线；禁用大面积圆形泡泡叠层。"),
    ])
    add_h2(doc, "数据识别原则")
    add_matrix(doc, ["对象", "主视觉", "辅助", "品牌化方式"], [
        ("完整比分", "比分 + 双方姓名", "阶段、日期、项目", "Ace Lime 强调胜方关键数字"),
        ("排名", "名次 + 球员 + 积分", "变化、场次、样本", "浅底为主，紫色硬阴影"),
        ("近期状态", "W / L 序列", "对手、比分、日期", "绿 / 红语义色 + 文字"),
        ("H2H", "交锋次数 + 胜场", "比赛清单", "左右对称，中轴 VS"),
    ], [1600, 3000, 2600, 2160])
    add_callout(doc, "EDITORIAL SCOREBOARD", "让比分成为海报，不让装饰替比分说话", "赛果视觉优先采用矩形记分牌、网格和清晰层级。Pride 彩虹只作为一条运动轨迹或边缘切面；冠军高光使用 Ace Lime，而不是再叠一层发光圆。", INK, LIME, WHITE)

    # 08 — Applications and governance
    add_page_break(doc)
    add_kicker(doc, "07 · APPLICATIONS & GOVERNANCE")
    add_h1(doc, "落地应用与品牌治理")
    add_h2(doc, "应用矩阵")
    add_matrix(doc, ["触点", "主识别", "建议规格", "关键要求"], [
        ("网站 / 小程序", "字标、色彩、数据层级", "Web / 375px", "双端同品牌，不复制两套 VI"),
        ("赛事海报", "比分、双方、日期、Spin 轨迹", "4:5 / 9:16 / A3", "比分优先，禁用泡泡叠层"),
        ("社交媒体", "抓马标题 + 真实人物", "1:1 / 4:5", "语气有趣但不羞辱成员"),
        ("服装正面", "HSAY + 六色轨迹", "左胸或居中", "保持字标安全空间"),
        ("服装背面", "今天演了，下次横扫", "大面积排字", "不与号码、赞助信息冲突"),
        ("奖牌 / 贴纸", "HSAY / H-S-A-Y 单字", "小尺寸", "使用单色简化版"),
    ], [1600, 2800, 1900, 3060])
    add_h2(doc, "文案语气")
    add_label_detail(doc, [
        ("赛前", "自信、邀请、带一点挑衅：缺你不可，等你上场。"),
        ("赛果", "先给事实，再给情绪：D-I-Y 6–3 登顶，今晚不演。"),
        ("失利", "不羞辱、不甩锅：今天演了，下次横扫。"),
        ("社群", "称呼亲近但尊重边界；避免替成员定义身份。"),
    ])
    add_h2(doc, "发布前 VI 检查")
    add_matrix(doc, ["检查项", "通过标准"], [
        ("字标", "标准写法、比例正确、安全空间完整"),
        ("色彩", "主色配比克制；胜负语义色没有混用"),
        ("图形与数据", "彩虹只有一条主轨迹；无装饰泡泡；比分与姓名优先可读"),
        ("影像与语气", "已授权、真实、多样；有锋芒、有温度，不冒犯、不排斥"),
    ], [2400, 6960])
    trailing = doc.paragraphs[-1]
    if not trailing.text:
        trailing._element.getparent().remove(trailing._element)

    core = doc.core_properties
    core.title = "HSAY 品牌视觉识别系统 VI 规范"
    core.subject = "HSAY Visual Identity Guidelines"
    core.author = "HSAY Brand Team"
    core.keywords = "HSAY, Tennis, Brand, VI, Visual Identity, Pride"
    core.comments = "Version 1.0"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
