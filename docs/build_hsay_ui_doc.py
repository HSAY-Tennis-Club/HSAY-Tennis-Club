from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "HSAY_UI设计规范_v1.0.docx"

INK = "141414"
PAPER = "F4F1EA"
WHITE = "FFFDF8"
LIME = "D7FF38"
VIOLET = "8263FF"
CORAL = "FF654F"
CYAN = "6EE6F2"
PINK = "FF78B7"
GREEN = "087958"
RED = "DC5D49"
MUTED = "716E68"
LINE = "D8D3CA"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, size=10.5, color=INK, bold=False, italic=False, ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(paragraph, before=0, after=6, line=1.25, keep=False):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False, before=0, after=6,
             line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    set_para(p, before, after, line, keep)
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size, color, bold, italic)
    return p


def add_kicker(doc, text):
    p = add_para(doc, text.upper(), 8.5, CORAL, True, before=0, after=5, line=1.0, keep=True)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    set_para(p, 18, 10, 1.0, True)
    set_run(p.add_run(text), 16, INK, True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    set_para(p, 14, 7, 1.0, True)
    set_run(p.add_run(text), 13, INK, True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    set_para(p, 10, 5, 1.0, True)
    set_run(p.add_run(text), 12, INK, True)
    return p


def add_label_detail(doc, rows, label_fill=PAPER):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, detail in rows:
        cells = table.add_row().cells
        shade(cells[0], label_fill)
        p0 = cells[0].paragraphs[0]
        set_para(p0, 0, 0, 1.15)
        set_run(p0.add_run(label), 9.5, INK, True)
        p1 = cells[1].paragraphs[0]
        set_para(p1, 0, 0, 1.2)
        set_run(p1.add_run(detail), 9.5, INK)
    set_table_geometry(table, [1700, 7660])
    add_para(doc, "", after=2)
    return table


def add_matrix(doc, headers, rows, widths, header_fill=INK, header_color=WHITE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.0)
        set_run(p.add_run(header), 8.5, header_color, True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            set_para(p, 0, 0, 1.15)
            color = GREEN if value in {"公开", "W", "胜"} else RED if value in {"登录后", "L", "负"} else INK
            set_run(p.add_run(str(value)), 8.8, color, value in {"公开", "登录后", "W", "L", "胜", "负"})
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, label, title, body, fill=INK, accent=LIME, text_color=WHITE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    set_para(p, 0, 4, 1.0)
    set_run(p.add_run(label.upper()), 8, accent, True)
    p2 = cell.add_paragraph()
    set_para(p2, 0, 5, 1.05)
    set_run(p2.add_run(title), 16, text_color, True)
    p3 = cell.add_paragraph()
    set_para(p3, 0, 0, 1.25)
    set_run(p3.add_run(body), 9.5, text_color)
    set_table_geometry(table, [9360])
    add_para(doc, "", after=3)
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, 8, MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 12, 10, 5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    set_para(hp, 0, 0, 1.0)
    set_run(hp.add_run("HSAY  /  PRODUCT UI SPECIFICATION"), 8, MUTED, True)
    footer = section.footer
    fp = footer.paragraphs[0]
    set_para(fp, 0, 0, 1.0)
    add_page_number(fp)
    return doc


def build():
    doc = setup_document()

    # Cover — editorial_cover with HSAY brand overrides.
    add_para(doc, "HSAY · HIT / SPIN / ACE / YOU", 9, CORAL, True, after=70, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "HSAY", 34, INK, True, after=3, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    add_para(doc, "网站与微信小程序 UI 设计规范", 22, INK, True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.05)
    add_para(doc, "场下是宝贝，场上撕得飞。", 14, VIOLET, True, after=34, align=WD_ALIGN_PARAGRAPH.CENTER)
    brand = doc.add_table(rows=1, cols=6)
    for cell, color in zip(brand.rows[0].cells, [CORAL, "FF9C2A", LIME, "44DFAC", "4DA6FF", VIOLET]):
        shade(cell, color)
        cell.text = ""
    set_table_geometry(brand, [1560] * 6, indent=0)
    add_para(doc, "", after=28)
    add_callout(doc, "VERSION 1.0 · 2026.08.24", "双端一致的数据体验，不同的交互外壳", "Web 保持宽屏信息效率；小程序按 375 CSS px / 750rpx 基准呈现，强调触控导航、紧凑纵向节奏和一手可读性。", PAPER, VIOLET, INK)
    add_para(doc, "适用范围：公开首页、赛事、完整赛果、积分与排名、球员档案、H2H、会员数据舱及密友备注。", 9.5, MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_kicker(doc, "01 · PRODUCT FRAME")
    add_h1(doc, "产品目标与访问层级")
    add_para(doc, "HSAY 是服务上海网球社群的赛事与球员数据产品。公开层用于展示俱乐部活力与竞技记录，登录层承载只对本人和授权关系开放的分析与备注。")
    add_callout(doc, "CORE PROMISE", "场下是宝贝，场上撕得飞。", "界面既要有硬核竞技的清晰与可信，也要保留 HSAY 的社群温度、Pride 包容性和适度抓马感。")
    add_h2(doc, "权限矩阵")
    add_matrix(doc, ["模块", "未登录", "登录后", "隐私说明"], [
        ("赛事日历 / 赛事清单", "公开", "公开", "展示时间、地点、赛制与状态"),
        ("完整赛果 / H2H", "公开", "公开", "只使用已登记公开比赛"),
        ("积分与排名 / 球员档案", "公开", "公开", "不显示联系方式与身份信息"),
        ("六维表现 / 近期状态", "登录后", "登录后", "仅本人及获授权管理员"),
        ("密友备注", "登录后", "登录后", "仅本人和被授权密友"),
    ], [2200, 1200, 1200, 4760])
    add_h2(doc, "关键体验原则")
    add_label_detail(doc, [
        ("数据先于装饰", "比分、胜负、排名、变化、样本量必须优先可读。"),
        ("公开与私密可感知", "登录按钮、锁定标签和说明文案必须让用户明确知道访问边界。"),
        ("品牌有性格", "用荧光绿、紫色与粗体制造识别度，但避免大面积纯黑造成压迫。"),
        ("双端同一事实源", "Web 与小程序共享赛事、球员和排名数据，只切换布局密度与导航。"),
    ])

    add_page_break(doc)
    add_kicker(doc, "02 · INFORMATION ARCHITECTURE")
    add_h1(doc, "信息架构与主导航")
    add_matrix(doc, ["层级", "入口", "核心任务", "默认可见性"], [
        ("首页", "Logo / 首页 Tab", "理解品牌、查看关键数据与赛事入口", "公开"),
        ("赛事", "赛事 Tab", "查日历、赛事安排、阶段与完整比分", "公开"),
        ("排名", "排名 Tab", "切换年度积分、单打实力、双打实力", "公开"),
        ("球员", "球员 Tab", "搜索球员、查看履历与公开统计", "公开"),
        ("H2H", "H2H 入口", "选择一对一或多人组合并查询历史交锋", "公开"),
        ("我的", "会员登录 / 我的 Tab", "查看六维表现、近期状态与密友备注", "登录后"),
    ], [1000, 1800, 4360, 2200])
    add_h2(doc, "Web 与小程序的结构关系")
    add_matrix(doc, ["维度", "Web", "小程序"], [
        ("宽度", "响应式宽屏；桌面优先利用横向空间", "固定 375 CSS px，窄屏时最大 100%"),
        ("导航", "顶部主导航 + 会员登录", "紧凑顶部栏 + 固定底部五栏 Tab"),
        ("布局", "两栏、四栏和宽表格并用", "单列为主，信息卡纵向连续"),
        ("节奏", "区块间距舒展，适合浏览与对比", "减少纵向留白，提高每屏信息量"),
        ("数据", "同一赛事、球员、排名和权限规则", "同一事实源，不创建平行数据"),
    ], [1600, 3880, 3880])
    add_callout(doc, "MODE SWITCH", "切换的是交互外壳，不是内容版本", "“Web / 小程序”按钮用于原型评审。生产小程序应复用相同的信息架构与数据接口，并采用微信端原生导航与安全区处理。", VIOLET, LIME, WHITE)

    add_page_break(doc)
    add_kicker(doc, "03 · MINI PROGRAM")
    add_h1(doc, "375px 小程序布局规范")
    add_label_detail(doc, [
        ("设计基准", "375 CSS px = 750rpx；左右内容边距 16px。"),
        ("窄屏回退", "视口小于 375px 时宽度为 100%，禁止横向滚动。"),
        ("顶部栏", "64px 高；保留 HSAY 与模式切换，隐藏桌面主导航及会员按钮。"),
        ("底部 Tab", "65px 高；首页、赛事、排名、球员、我的；固定于安全区上方。"),
        ("触控目标", "主要按钮最小 44px；列表行最小 66px；选择控件不小于 40px。"),
    ])
    add_h2(doc, "紧凑化节奏")
    add_matrix(doc, ["区域", "原密度", "新密度", "设计理由"], [
        ("首屏", "上 66 / 下 48px", "上 42 / 下 30px", "标题与 CTA 更快进入首屏"),
        ("通用区块", "上下 78px", "上下 54px", "减少连续长页滚动距离"),
        ("赛事卡", "最小 125px", "最小 108px", "保留日期块与状态，压缩空白"),
        ("比分行", "最小 76px", "最小 66px", "姓名与比分仍保持双行可读"),
        ("排名行", "最小 90px", "最小 76px", "保留名次、球员、积分与变化"),
        ("球员卡", "最小 108px", "最小 92px", "头像缩至 50px，信息不丢失"),
        ("品牌页脚", "最小 430px", "最小 330px", "保留宣言，减少末尾空段"),
    ], [1450, 1700, 1700, 4510])
    add_h2(doc, "首屏顺序")
    add_label_detail(doc, [
        ("01", "品牌眉题：SHANGHAI · 150 PLAYERS · EST. 2024"),
        ("02", "主标题：场下是宝贝，场上撕得飞。"),
        ("03", "赛事 CTA + 赛季排名次入口"),
        ("04", "冠军女性杯赛果卡，随后进入俱乐部数据条"),
    ], label_fill=LIME)

    add_kicker(doc, "04 · PUBLIC MODULES")
    add_h1(doc, "公开模块组件规范")
    add_h2(doc, "赛事日历与赛事清单")
    add_para(doc, "赛事条目由日期块、状态、赛事名称、场地、赛制摘要和详情入口组成。已结束、进行中、待公布使用不同语义色；无日期时显示“—”，不猜测时间。")
    add_h2(doc, "完整赛果")
    add_label_detail(doc, [
        ("阶段切换", "决赛、半决赛 A、半决赛 B、三四名决赛。"),
        ("比分头部", "赛事阶段、队名、总比分、场数和计分规则。"),
        ("逐场赛果", "左方 / 比分 / 右方；名字大于辅助标签；单双打明确标注。"),
        ("胜负颜色", "WIN 使用深绿 #087958；LOST 使用砖红 #DC5D49。"),
        ("图例", "记分卡底部必须保留 W/L 颜色指示，不能只依赖颜色判断。"),
    ])
    add_h2(doc, "积分与排名")
    add_matrix(doc, ["排名类型", "核心数值", "辅助信息", "异常处理"], [
        ("年度积分", "PTS / 排名", "历史最佳、总计站次、本期站次、变化", "缺少变化显示 —"),
        ("单打实力", "Elo", "样本场次、稳定性标签", "数据未导入时显示回算说明"),
        ("双打实力", "Elo Beta", "双打场次、变化", "明确 Beta 与估算口径"),
    ], [1800, 1800, 3560, 2200])
    add_h2(doc, "球员与 H2H")
    add_para(doc, "球员清单支持名字搜索；公开卡片展示年度排名、已登记比赛、冠军数或积分。H2H 支持左右球员选择和交换，展示交锋次数、胜场、领先方及相关比赛记录。")

    add_kicker(doc, "05 · MEMBER EXPERIENCE")
    add_h1(doc, "会员数据舱与近期状态")
    add_callout(doc, "RECENT FORM", "近期状态必须来自近期比赛输赢", "组件不再使用难以解释的抽象柱状曲线，而是展示最近 8 场的 W/L 序列、阶段统计和最近 4 场明细。")
    add_h2(doc, "近期状态组件")
    add_matrix(doc, ["顺序", "结果", "对手", "项目", "比分", "日期"], [
        ("1", "W", "夏和雪", "双打", "15–8", "08.21"),
        ("2", "W", "Ivan", "双打", "15–11", "08.18"),
        ("3", "L", "宇凡", "单打", "11–15", "08.12"),
        ("4", "W", "CY", "双打", "15–9", "08.06"),
        ("5", "L", "Loker", "单打", "8–15", "07.28"),
        ("6", "W", "猪猪", "双打", "15–6", "07.21"),
        ("7", "W", "川林贯空", "双打", "15–13", "07.14"),
        ("8", "L", "刀刀", "单打", "12–15", "07.08"),
    ], [700, 700, 2240, 1200, 1900, 2620])
    add_para(doc, "说明：以上为 UI 原型示例数据；接入正式比赛库后，应按比赛日期倒序自动生成。", 8.5, MUTED, italic=True, before=2, after=8)
    add_h2(doc, "六维分析与备注")
    add_matrix(doc, ["模块", "当前 UI", "产品约束"], [
        ("六维表现", "实力、稳定、压制、韧性、调整、关键分", "维度定义与权重未定稿；必须显示样本量"),
        ("雷达图", "按 0–100 数值实时渲染", "不得使用固定形状；数值变更必须改变多边形"),
        ("密友备注", "训练建议、标签、日期", "仅本人和被授权密友可见"),
        ("公开摘要", "排名、积分、参赛与公开胜率", "不得泄露私密指标及备注"),
    ], [1800, 3100, 4460])

    add_page_break(doc)
    add_kicker(doc, "06 · VISUAL SYSTEM")
    add_h1(doc, "视觉系统与识别规则")
    add_h2(doc, "品牌色")
    add_matrix(doc, ["角色", "色值", "主要用途", "避免事项"], [
        ("Ink", "#141414", "正文、边框、深色卡片", "排名区避免大面积铺满"),
        ("Paper", "#F4F1EA", "页面底色、柔和留白", "不要与纯白卡片失去层级"),
        ("Lime", "#D7FF38", "Ace、高光、主要 CTA", "不用于长篇正文"),
        ("Violet", "#8263FF", "会员分析、强调阴影", "控制面积，保证白字对比"),
        ("Coral", "#FF654F", "状态、眉题、抓马强调", "不替代 LOST 的语义红"),
        ("Win", "#087958", "胜方、W、正向状态", "必须配合 WIN / W 文本"),
        ("Lost", "#DC5D49", "负方、L、负向状态", "必须配合 LOST / L 文本"),
    ], [1400, 1400, 3400, 3160])
    add_h2(doc, "字体与层级")
    add_label_detail(doc, [
        ("展示标题", "超粗无衬线；中文优先 PingFang SC / Microsoft YaHei；紧字距。"),
        ("正文", "11–14px（Web）/ 10–13px（小程序）；颜色为 Ink 或 Muted。"),
        ("数据", "等宽字或视觉等宽；比分、积分、排名必须比标签更突出。"),
        ("辅助标签", "7–10px 大写英文或短中文；提高字距，避免承担核心信息。"),
    ])
    add_h2(doc, "卡片语言")
    add_para(doc, "以 1px 深色边框、轻微硬阴影、荧光色重点和少量圆角构成。赛事与排名优先使用浅底，会员画像可使用深底，但每屏最多一个大面积深色焦点。")

    add_page_break(doc)
    add_kicker(doc, "07 · INTERACTION & QA")
    add_h1(doc, "交互、响应式与验收标准")
    add_h2(doc, "交互状态")
    add_matrix(doc, ["控件", "默认", "激活 / 成功", "空 / 错误"], [
        ("Web / 小程序切换", "Web 高亮", "小程序高亮并出现状态提示", "保持当前模式，不闪烁"),
        ("排名类型", "年度积分", "激活项使用高识别底色", "未导入数据展示说明"),
        ("赛事阶段", "决赛", "切换后比分卡整体更新", "无对阵时展示空状态"),
        ("H2H 选择", "默认两位球员", "交换后胜负同步反转", "无交锋时不推断胜者"),
    ], [1900, 2200, 3060, 2200])
    add_h2(doc, "无障碍与可读性")
    add_label_detail(doc, [
        ("颜色", "W/L、升降、锁定状态都必须同时提供文字或图形标签。"),
        ("键盘", "按钮、Tab、选择器、搜索框和链接均可获得焦点并操作。"),
        ("触控", "核心交互目标不小于 44px；底部 Tab 具有明确当前态。"),
        ("数据可信", "Elo、积分、胜率和画像必须展示计算口径或样本说明。"),
    ])
    add_h2(doc, "发布验收")
    add_matrix(doc, ["检查项", "通过标准"], [
        ("375px 小程序模式", "主体与底栏同宽；窄于 375px 自动收缩；无横向溢出"),
        ("紧凑度", "首屏、区块、赛事卡、比分行和排名行使用新版间距"),
        ("近期状态", "W/L 序列、5胜3负摘要、最近四场明细和颜色图例同时存在"),
        ("隐私", "未登录不可访问会员数据；密友备注不出现在公开层"),
    ], [2500, 6860])
    core = doc.core_properties
    core.title = "HSAY 网站与微信小程序 UI 设计规范"
    core.subject = "HSAY Web / Mini Program UI specification"
    core.author = "HSAY Product Team"
    core.keywords = "HSAY, Tennis, UI, WeChat Mini Program, Web"
    core.comments = "Version 1.0"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
